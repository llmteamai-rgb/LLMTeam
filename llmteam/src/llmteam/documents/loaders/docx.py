"""
DOCX loader for Microsoft Word files.

Requires: python-docx (pip install llmteam-ai[docx])
"""

from pathlib import Path
from typing import List, Union

from llmteam.documents.models import Document
from llmteam.documents.loaders.base import BaseLoader


class DocxLoader(BaseLoader):
    """
    Loader for Microsoft Word (.docx) files.

    Requires the python-docx package:
        pip install python-docx
        # or
        pip install llmteam-ai[docx]
    """

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self, paragraph_separator: str = "\n\n"):
        """
        Initialize the DOCX loader.

        Args:
            paragraph_separator: String to separate paragraphs.
        """
        self.paragraph_separator = paragraph_separator
        self._ensure_docx()

    def _ensure_docx(self) -> None:
        """Ensure python-docx is installed."""
        try:
            import docx  # noqa: F401
        except ImportError:
            raise ImportError(
                "DOCX support requires 'python-docx'. Install it with:\n"
                "  pip install python-docx\n"
                "or:\n"
                "  pip install llmteam-ai[docx]"
            )

    def load(self, source: Union[str, Path]) -> List[Document]:
        """
        Load a DOCX file as a Document.

        Args:
            source: Path to the DOCX file.

        Returns:
            List containing a single Document with the file content.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file extension is not supported.
        """
        import docx

        path = Path(source)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        if not self.supports(path):
            raise ValueError(f"Unsupported file extension: {path.suffix}. Expected: .docx")

        doc = docx.Document(str(path))

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = self.paragraph_separator.join(paragraphs)

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        if table_texts:
            table_content = "\n".join(table_texts)
            if content:
                content = content + self.paragraph_separator + table_content
            else:
                content = table_content

        return [
            Document(
                content=content,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "extension": ".docx",
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables),
                    "size_bytes": path.stat().st_size,
                },
            )
        ]

    def supports(self, path: Union[str, Path]) -> bool:
        """Check if this loader supports the given file."""
        return self._get_extension(path) in self.SUPPORTED_EXTENSIONS


__all__ = ["DocxLoader"]
